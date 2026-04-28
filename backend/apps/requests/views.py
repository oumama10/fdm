from django.utils import timezone
from django.http import HttpResponse
from rest_framework import permissions, status, viewsets
from rest_framework.decorators import action
from rest_framework.response import Response

from apps.core.permissions import (
    IsChefService,
    IsChefServiceOwner,
    IsGestionnaireOrAdmin,
)

from .models import Demande
from .serializers import DemandeCreateSerializer, DemandeSerializer


REQUESTER_OPTIONS_BY_TYPE = {
    "service": {
        "beneficiaire_types": ["personnel", "unite"],
        "beneficiaires": [
            "Chef de service",
            "Secretaire",
            "Personnel du service",
            "Unite du service",
        ],
    },
    "chu": {
        "beneficiaire_types": ["personnel", "lieu"],
        "beneficiaires": ["Chef", "Secretaire", "PA", "Salle de cours"],
    },
    "decanat": {
        "beneficiaire_types": ["personnel"],
        "beneficiaires": [
            "Vice doyen pedagogie",
            "Vice doyen recherche",
            "Vice doyen pharmacie",
            "Vice doyen dentaire",
            "Secretaire general",
            "Personnel dedie",
        ],
    },
    "pharmacie": {
        "beneficiaire_types": ["personnel"],
        "beneficiaires": ["Chef departement", "Personnel departement"],
    },
    "dentaire": {
        "beneficiaire_types": ["personnel"],
        "beneficiaires": ["Chef departement", "Personnel departement"],
    },
    "labo": {
        "beneficiaire_types": ["personnel"],
        "beneficiaires": ["Directeur labo", "Personnel labo"],
    },
    "association": {
        "beneficiaire_types": ["personnel"],
        "beneficiaires": ["President association"],
    },
}

ONLY_ALLOWED_DEMANDEUR = "Chef de service"


def _map_service_type_to_requester_type(service_type: str) -> str:
    if service_type == "administratif":
        return "service"
    return service_type


class DemandeViewSet(viewsets.ModelViewSet):
    """
    Permissions per action
    ----------------------
    create                : IsChefService
    list                  : IsGestionnaireOrAdmin → all
                            IsChefService        → own only
    retrieve              : IsGestionnaireOrAdmin → any
                            IsChefService        → own only (object-level)
    update/partial_update : IsGestionnaireOrAdmin
    destroy               : IsGestionnaireOrAdmin

    Custom actions
    --------------
    POST /demandes/{id}/valider/  — IsGestionnaireOrAdmin
    POST /demandes/{id}/refuser/  — IsGestionnaireOrAdmin
    GET  /demandes/{id}/download_pdf/ — IsGestionnaireOrAdmin | owner
    POST /demandes/{id}/creer_commande_interne/ — IsGestionnaireOrAdmin
    POST /demandes/{id}/signer_commande_interne/ — IsGestionnaireOrAdmin
    """

    http_method_names = ["get", "post", "put", "patch", "delete", "head", "options"]

    def get_serializer_class(self):
        if self.action == "create":
            return DemandeCreateSerializer
        return DemandeSerializer

    def get_permissions(self):
        if self.action == "requester_options":
            return [permissions.IsAuthenticated()]
        if self.action == "create":
            return [IsChefService()]
        if self.action in (
            "update",
            "partial_update",
            "destroy",
            "valider",
            "refuser",
            "creer_commande_interne",
            "signer_commande_interne",
        ):
            return [IsGestionnaireOrAdmin()]
        # list + retrieve: both roles allowed — queryset/object gate narrows access
        return [(IsGestionnaireOrAdmin | IsChefService)()]

    def get_queryset(self):
        qs = Demande.objects.select_related(
            "id_chef_demandeur",
            "id_service",
            "id_valide_par",
        ).prefetch_related(
            "lignes__id_ressource__id_categorie",
            "lignes__id_ressource__id_sous_categorie__id_parent_sous_categorie",
        )

        user = self.request.user
        # Chefs de service see only their own demandes
        if (
            user.is_authenticated
            and user.id_role
            and user.id_role.nom_role == "chef_service"
        ):
            qs = qs.filter(id_chef_demandeur=user)
        return qs

    def get_object(self):
        obj = super().get_object()
        user = self.request.user
        # Object-level guard for chef_service on retrieve
        if (
            self.action in ("retrieve", "download_pdf")
            and user.is_authenticated
            and user.id_role
            and user.id_role.nom_role == "chef_service"
        ):
            self.check_object_permissions(self.request, obj)
            # IsChefServiceOwner is checked explicitly here
            perm = IsChefServiceOwner()
            if not perm.has_object_permission(self.request, self, obj):
                self.permission_denied(self.request)
        return obj

    def perform_create(self, serializer):
        serializer.save(id_chef_demandeur=self.request.user)

    # ── helpers ──────────────────────────────────────────────────────────────

    @staticmethod
    def _notify_chef(demande: Demande, titre: str, message: str) -> None:
        """
        Send a web Notification to the chef who submitted the demande.
        Silently swallows all exceptions — notifications must not block the
        main action.
        """
        try:
            from apps.alerts.models import Notification  # noqa: PLC0415
            from django.contrib.contenttypes.models import ContentType  # noqa: PLC0415

            Notification.objects.create(
                id_destinataire=demande.id_chef_demandeur,
                type_notification="validation_requise",
                titre=titre,
                message=message,
                canal="web",
                content_type=ContentType.objects.get_for_model(Demande),
                object_id=demande.pk,
            )
        except Exception:
            pass

    # ── custom actions ────────────────────────────────────────────────────────

    @action(detail=False, methods=["get"], url_path="requester-options")
    def requester_options(self, request):
        from apps.users.models import Service  # noqa: PLC0415

        service_id = request.query_params.get("id_service")
        if not service_id:
            return Response(
                {"detail": "Le parametre id_service est obligatoire."},
                status=status.HTTP_400_BAD_REQUEST,
            )

        service = Service.objects.filter(pk=service_id).first()
        if not service:
            return Response(
                {"detail": "Service introuvable."},
                status=status.HTTP_404_NOT_FOUND,
            )

        requester_type = _map_service_type_to_requester_type(service.type_service)
        config = REQUESTER_OPTIONS_BY_TYPE.get(
            requester_type,
            {"beneficiaire_types": ["personnel"], "beneficiaires": []},
        )

        return Response(
            {
                "id_service": service.id_service,
                "nom_service": service.nom_service,
                "type_demandeur": requester_type,
                "beneficiaire_types": config["beneficiaire_types"],
                "beneficiaires": [ONLY_ALLOWED_DEMANDEUR],
            },
            status=status.HTTP_200_OK,
        )

    @action(detail=True, methods=["post"], url_path="valider")
    def valider(self, request, pk=None):
        from django.db import transaction  # noqa: PLC0415
        from apps.decharge.models import Decharge, LigneDecharge, SignatureDecharge  # noqa: PLC0415
        from apps.resources.models import InstanceRessource, Stock  # noqa: PLC0415

        demande = self.get_object()

        if demande.statut != "en_cours":
            return Response(
                {"detail": f"Impossible de valider une demande au statut '{demande.statut}'."},
                status=status.HTTP_400_BAD_REQUEST,
            )

        lignes_qs = demande.lignes.select_related("id_ressource__id_categorie").all()
        lignes = list(lignes_qs)
        if not lignes:
            return Response(
                {"detail": "La demande ne contient aucune ligne."},
                status=status.HTTP_400_BAD_REQUEST,
            )

        # Build lookup: id_ligne → payload entry
        payload_map = {
            int(item["id_ligne"]): item
            for item in request.data.get("lignes", [])
            if "id_ligne" in item
        }

        decharge_lignes_data = []

        with transaction.atomic():
            for ligne in lignes:
                ressource = ligne.id_ressource
                ld = payload_map.get(ligne.id_ligne, {})

                if ressource.is_consommable:
                    # ── Consommable: deduct from stock ──────────────────────
                    requested = max(0, min(int(ld.get("quantite_accordee", 0)), ligne.quantite_demandee))
                    try:
                        stock = Stock.objects.select_for_update().get(id_ressource=ressource)
                        qty = min(requested, max(stock.quantite_disponible, 0))
                        stock.quantite_disponible -= qty
                        stock.save(update_fields=["quantite_disponible"])
                    except Stock.DoesNotExist:
                        qty = 0

                    ligne.quantite_accordee = qty
                    ligne.save(update_fields=["quantite_accordee"])

                    if qty > 0:
                        decharge_lignes_data.append({
                            "ressource": ressource,
                            "quantite": qty,
                            "type_ligne": "consommable",
                            "instance": None,
                        })

                else:
                    # ── Bien inventaire: assign selected instances ──────────
                    raw_ids = [int(i) for i in ld.get("instances", []) if str(i).isdigit()]
                    raw_ids = raw_ids[: ligne.quantite_demandee]

                    valid_instances = list(
                        InstanceRessource.objects.select_for_update().filter(
                            pk__in=raw_ids,
                            id_ressource=ressource,
                            statut="en_stock",
                        )
                    )
                    if valid_instances:
                        InstanceRessource.objects.filter(
                            pk__in=[inst.pk for inst in valid_instances]
                        ).update(statut="en_service", id_service_actuel=demande.id_service)

                    qty = len(valid_instances)
                    ligne.quantite_accordee = qty
                    ligne.save(update_fields=["quantite_accordee"])

                    for inst in valid_instances:
                        decharge_lignes_data.append({
                            "ressource": ressource,
                            "quantite": 1,
                            "type_ligne": "bien_inventaire",
                            "instance": inst,
                        })

            # ── Compute global status ───────────────────────────────────────
            total = len(lignes)
            fully_served = sum(1 for l in lignes if l.quantite_accordee >= l.quantite_demandee)
            total_accordee = sum(l.quantite_accordee for l in lignes)

            if total_accordee == 0:
                new_statut = "refusee"
            elif fully_served == total:
                new_statut = "totale"
            else:
                new_statut = "partielle"

            demande.statut = new_statut
            demande.id_valide_par = request.user
            demande.date_validation = timezone.now()
            demande.save(update_fields=["statut", "id_valide_par_id", "date_validation"])

            # ── Auto-create Decharge (once) ─────────────────────────────────
            existing = Decharge.objects.filter(id_demande=demande).first()
            if not existing and decharge_lignes_data:
                decharge = Decharge.objects.create(
                    id_demande=demande,
                    id_genere_par=request.user,
                    date_livraison=timezone.localdate(),
                )
                SignatureDecharge.objects.create(
                    id_decharge=decharge,
                    id_chef_service=demande.id_chef_demandeur,
                    statut="en_attente",
                )
                for dl in decharge_lignes_data:
                    LigneDecharge.objects.create(
                        id_decharge=decharge,
                        id_ressource=dl["ressource"],
                        quantite=dl["quantite"],
                        type_ligne=dl["type_ligne"],
                        id_instance_ressource=dl["instance"],
                    )

        self._notify_chef(
            demande,
            titre="Demande traitée",
            message=f"Votre demande #{demande.id_demande} est {demande.statut}.",
        )

        return Response(DemandeSerializer(demande).data, status=status.HTTP_200_OK)

    @action(detail=True, methods=["get"], url_path="download_pdf")
    def download_pdf(self, request, pk=None):
        import subprocess
        import tempfile
        from pathlib import Path

        import openpyxl

        demande = self.get_object()

        with tempfile.TemporaryDirectory() as tmp_dir:
            tmp_path = Path(tmp_dir)
            xlsx_path = tmp_path / f"demande_{demande.id_demande}.xlsx"

            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "Demande"

            ws["A1"] = "Demande"
            ws["A2"] = f"Référence: DEM-{demande.id_demande:06d}"
            ws["A3"] = f"Date: {demande.date_demande.strftime('%d/%m/%Y %H:%M')}"
            ws["A4"] = f"Demandeur: {demande.id_chef_demandeur.nom_complet if demande.id_chef_demandeur else '—'}"
            ws["A5"] = f"Service: {demande.id_service.nom_service if demande.id_service else '—'}"
            ws["A6"] = f"Urgence: {demande.urgence}"
            ws["A7"] = f"Bénéficiaire: {demande.beneficiaire_nom or '—'} ({demande.beneficiaire_type or '—'})"
            ws["A8"] = f"Détail bénéficiaire: {demande.beneficiaire_detail or '—'}"
            ws["A10"] = "Justification"
            ws["A11"] = demande.justification or "—"

            ws["A13"] = "Article"
            ws["B13"] = "Qté demandée"
            ws["C13"] = "Qté accordée"
            ws["D13"] = "Disponibilité %"

            row = 14
            for ligne in demande.lignes.select_related("id_ressource").all():
                ws.cell(row=row, column=1, value=ligne.id_ressource.designation if ligne.id_ressource else "—")
                ws.cell(row=row, column=2, value=ligne.quantite_demandee)
                ws.cell(row=row, column=3, value=ligne.quantite_accordee)
                ws.cell(row=row, column=4, value=ligne.disponibilite_pct)
                row += 1

            wb.save(str(xlsx_path))
            wb.close()

            result = subprocess.run(
                [
                    "libreoffice",
                    "--headless",
                    "--convert-to",
                    "pdf",
                    "--outdir",
                    str(tmp_path),
                    str(xlsx_path),
                ],
                capture_output=True,
                text=True,
                timeout=60,
                check=False,
            )

            if result.returncode != 0:
                return Response(
                    {"detail": f"Conversion PDF impossible: {result.stderr.strip()}"},
                    status=status.HTTP_500_INTERNAL_SERVER_ERROR,
                )

            pdf_path = tmp_path / f"demande_{demande.id_demande}.pdf"
            if not pdf_path.is_file():
                return Response(
                    {"detail": "Le PDF n'a pas pu être généré."},
                    status=status.HTTP_500_INTERNAL_SERVER_ERROR,
                )

            pdf_bytes = pdf_path.read_bytes()

        response = HttpResponse(pdf_bytes, content_type="application/pdf")
        response["Content-Disposition"] = (
            f'attachment; filename="demande-{demande.id_demande}.pdf"'
        )
        return response

    @action(detail=True, methods=["get"], url_path="imprimer")
    def imprimer(self, request, pk=None):
        """
        Generate an institutional Fiche de Décharge PDF matching the FMPDF template layout.
        Drop logo_header.jpg (or .png) in backend/apps/decharge/templates/ for the logo.
        """
        import os
        import subprocess
        import tempfile
        from pathlib import Path

        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Cm, Inches, Pt

        # Robust helper: find child or create it (avoids get_or_add_* version issues)
        def _child(parent_el, w_tag):
            found = parent_el.find(qn(w_tag))
            if found is None:
                found = OxmlElement(w_tag)
                parent_el.append(found)
            return found

        def _set_tbl_borders(tbl, sz="12", color="000000", val="single"):
            tblPr = _child(tbl._tbl, "w:tblPr")
            for old in tblPr.findall(qn("w:tblBorders")):
                tblPr.remove(old)
            bdr = OxmlElement("w:tblBorders")
            for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
                el = OxmlElement(f"w:{side}")
                el.set(qn("w:val"), val)
                el.set(qn("w:sz"), sz)
                el.set(qn("w:space"), "0")
                el.set(qn("w:color"), color)
                bdr.append(el)
            tblPr.append(bdr)

        def _remove_tbl_borders(tbl):
            _set_tbl_borders(tbl, sz="0", color="auto", val="none")

        def _center_tbl(tbl, pct=60):
            tblPr = _child(tbl._tbl, "w:tblPr")
            for old in tblPr.findall(qn("w:tblW")):
                tblPr.remove(old)
            for old in tblPr.findall(qn("w:jc")):
                tblPr.remove(old)
            w = OxmlElement("w:tblW")
            w.set(qn("w:w"), str(pct * 50))
            w.set(qn("w:type"), "pct")
            tblPr.append(w)
            jc = OxmlElement("w:jc")
            jc.set(qn("w:val"), "center")
            tblPr.append(jc)

        def _cell_bg(cell, fill_hex):
            tcPr = _child(cell._tc, "w:tcPr")
            for old in tcPr.findall(qn("w:shd")):
                tcPr.remove(old)
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), fill_hex)
            tcPr.append(shd)

        def _cell_margins(cell, twips=120):
            tcPr = _child(cell._tc, "w:tcPr")
            for old in tcPr.findall(qn("w:tcMar")):
                tcPr.remove(old)
            mar = OxmlElement("w:tcMar")
            for side in ("top", "bottom", "left", "right"):
                m = OxmlElement(f"w:{side}")
                m.set(qn("w:w"), str(twips))
                m.set(qn("w:type"), "dxa")
                mar.append(m)
            tcPr.append(mar)

        def _bold_centered(cell, text, size=10):
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(text)
            r.bold = True
            r.font.size = Pt(size)

        demande = self.get_object()
        lignes = demande.lignes.select_related(
            "id_ressource__id_categorie",
            "id_ressource__id_sous_categorie__id_parent_sous_categorie",
        ).all()

        est_bien_inventaire = any(
            l.id_ressource
            and l.id_ressource.id_categorie
            and l.id_ressource.id_categorie.nom_categorie == "Bien Inventaire"
            for l in lignes
        )

        # Collect assigned instances from LigneDecharge (set during valider)
        instances_by_ressource = {}
        try:
            from apps.decharge.models import Decharge  # noqa: PLC0415

            decharge_obj = Decharge.objects.filter(id_demande=demande).first()
            if decharge_obj:
                for ld in decharge_obj.lignes.select_related("id_instance_ressource").filter(
                    type_ligne="bien_inventaire"
                ):
                    if ld.id_instance_ressource:
                        instances_by_ressource.setdefault(ld.id_ressource_id, []).append(
                            ld.id_instance_ressource
                        )
        except Exception:
            pass

        service_nom = demande.id_service.nom_service if demande.id_service else "—"
        date_str = demande.date_demande.strftime("%d/%m/%Y")

        # Title is static — the Word template (or title box) always says "DECHARGE"
        title_text = "DECHARGE"

        # ---- Build document ----
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

        # -- Logo header --
        templates_dir = Path(__file__).resolve().parent.parent / "decharge" / "templates"
        logo_path = next(
            (templates_dir / f for f in ("logo_header.png", "logo_header.jpg", "logo_header.jpeg")
             if (templates_dir / f).exists()),
            None,
        )
        if logo_path:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(str(logo_path), width=Inches(6.3))
        else:
            hdr = doc.add_table(rows=1, cols=3)
            _remove_tbl_borders(hdr)
            lc, mc, rc = hdr.cell(0, 0), hdr.cell(0, 1), hdr.cell(0, 2)
            lc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = lc.paragraphs[0].add_run(
                "FACULTÉ DE MÉDECINE, DE PHARMACIE\nET DE MÉDECINE DENTAIRE"
            )
            r.bold = True
            r.font.size = Pt(8)
            mc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            mc.paragraphs[0].add_run("[ LOGO ]").font.size = Pt(10)
            rc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            r2 = rc.paragraphs[0].add_run(
                "UNIVERSITÉ SIDI MOHAMED\nBEN ABDELLAH DE FES"
            )
            r2.bold = True
            r2.font.size = Pt(8)

        # Horizontal rule under header
        hr = doc.add_paragraph()
        hr_ppr = hr._p.get_or_add_pPr()
        hr_bdr = OxmlElement("w:pBdr")
        hr_bot = OxmlElement("w:bottom")
        hr_bot.set(qn("w:val"), "single")
        hr_bot.set(qn("w:sz"), "6")
        hr_bot.set(qn("w:space"), "1")
        hr_bot.set(qn("w:color"), "000000")
        hr_bdr.append(hr_bot)
        hr_ppr.append(hr_bdr)

        # -- "MAGASIN ET STOCK"  /  "Fès, le {date}" --
        lr = doc.add_table(rows=1, cols=2)
        _remove_tbl_borders(lr)
        lc, rc = lr.cell(0, 0), lr.cell(0, 1)
        lc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        rl = lc.paragraphs[0].add_run("MAGASIN ET STOCK")
        rl.underline = True
        rl.font.size = Pt(10)
        rc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rr = rc.paragraphs[0].add_run(f"Fès,  le {date_str}")
        rr.italic = True
        rr.underline = True
        rr.font.size = Pt(11)

        doc.add_paragraph()
        doc.add_paragraph()

        # -- Title box (gray-filled, bordered, centered) --
        tb = doc.add_table(rows=1, cols=1)
        _set_tbl_borders(tb, sz="18", color="808080")
        _center_tbl(tb, pct=55)
        tc = tb.cell(0, 0)
        _cell_bg(tc, "D3D3D3")
        _cell_margins(tc, twips=160)
        tp = tc.paragraphs[0]
        tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tr = tp.add_run(title_text)
        tr.bold = True
        tr.font.size = Pt(15)

        doc.add_paragraph()
        doc.add_paragraph()

        # -- Helper: format a list of inventory numbers into ranges --
        def _fmt_inv(numeros):
            import re as _re
            if not numeros:
                return "—"
            if len(numeros) == 1:
                return numeros[0]
            parsed = []
            for n in numeros:
                m = _re.match(r"^(.+)-(\d+)$", n)
                if m:
                    parsed.append((m.group(1), int(m.group(2)), n))
                else:
                    parsed.append((None, None, n))
            if any(p[0] is None for p in parsed):
                return ", ".join(numeros)
            prefixes = {p[0] for p in parsed}
            if len(prefixes) > 1:
                return ", ".join(numeros)
            prefix = parsed[0][0]
            sequences = sorted(p[1] for p in parsed)
            ranges, start, end = [], sequences[0], sequences[0]
            for seq in sequences[1:]:
                if seq == end + 1:
                    end = seq
                else:
                    ranges.append((start, end))
                    start = end = seq
            ranges.append((start, end))
            parts = []
            for s, e in ranges:
                if s == e:
                    parts.append(f"{prefix}-{str(s).zfill(4)}")
                else:
                    parts.append(f"{prefix}-{str(s).zfill(4)} à {prefix}-{str(e).zfill(4)}")
            return ", ".join(parts)

        # -- Articles table --
        if est_bien_inventaire:
            tbl = doc.add_table(rows=1, cols=4)
            _set_tbl_borders(tbl)
            hc = tbl.rows[0].cells
            _bold_centered(hc[0], "DÉSIGNATION DU MATÉRIEL")
            _bold_centered(hc[1], "N° INVENTAIRE")
            _bold_centered(hc[2], "QTÉ")
            _bold_centered(hc[3], "AFFECTATION")

            # Build lignes as dicts — affectation lives only inside each ligne,
            # never as a top-level variable (mirrors the docxtpl context shape).
            lignes_ctx = []
            for ligne in lignes:
                res = ligne.id_ressource
                if not res:
                    continue
                instances = instances_by_ressource.get(res.pk, [])
                if instances:
                    numeros = [i.numero_inventaire for i in instances if i.numero_inventaire]
                    lignes_ctx.append({
                        "designation": res.designation,
                        "n_inventaire": _fmt_inv(numeros),
                        "qte": f"{len(instances):02d}",
                        "affectation": service_nom,
                    })
                else:
                    lignes_ctx.append({
                        "designation": res.designation,
                        "n_inventaire": "—",
                        "qte": f"{ligne.quantite_demandee or 1:02d}",
                        "affectation": service_nom,
                    })

            for row_data in lignes_ctx:
                rc = tbl.add_row().cells
                rc[0].text = row_data["designation"]
                rc[1].text = row_data["n_inventaire"]
                rc[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                rc[2].text = row_data["qte"]
                rc[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                rc[3].text = row_data["affectation"]

            # Merge AFFECTATION column vertically — show lignes[0].affectation once,
            # spanning all data rows, matching the official template layout.
            if len(lignes_ctx) > 1:
                top_cell = tbl.rows[1].cells[3]
                bot_cell = tbl.rows[-1].cells[3]
                top_cell.merge(bot_cell)
                # Show affectation from the first ligne only (all share the same service)
                top_cell.text = lignes_ctx[0]["affectation"]
                top_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = top_cell.paragraphs[0].runs
                if r:
                    r[0].bold = True
        else:
            tbl = doc.add_table(rows=1, cols=2)
            _set_tbl_borders(tbl)
            hc = tbl.rows[0].cells
            _bold_centered(hc[0], "ARTICLE")
            _bold_centered(hc[1], "QUANTITÉ")

            for ligne in lignes:
                res = ligne.id_ressource
                if not res:
                    continue
                rc = tbl.add_row().cells
                rc[0].text = res.designation
                rc[1].text = str(ligne.quantite_accordee or ligne.quantite_demandee or 0)
                rc[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()
        doc.add_paragraph()

        sig = doc.add_paragraph()
        sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        sr = sig.add_run("SIGNE : CHEF DE SERVICE")
        sr.bold = True
        sr.font.size = Pt(11)

        import shutil
        import threading

        tmp_dir = tempfile.mkdtemp()
        try:
            docx_path = os.path.join(tmp_dir, f"decharge_{demande.id_demande}.docx")
            doc.save(docx_path)

            pdf_path = os.path.join(tmp_dir, f"decharge_{demande.id_demande}.pdf")

            # -- Convert DOCX → PDF --
            # Try LibreOffice first (no COM, no threading issues)
            import shlex
            lo_paths = [
                r"C:\Program Files\LibreOffice\program\soffice.exe",
                r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
                "soffice",
            ]
            soffice = next(
                (p for p in lo_paths if os.path.exists(p) if p != "soffice"), None
            ) or (shutil.which("soffice"))

            if soffice:
                proc = subprocess.run(
                    [soffice, "--headless", "--convert-to", "pdf",
                     "--outdir", tmp_dir, docx_path],
                    capture_output=True, timeout=90,
                )
                lo_out = os.path.join(
                    tmp_dir,
                    os.path.splitext(os.path.basename(docx_path))[0] + ".pdf",
                )
                if os.path.isfile(lo_out) and lo_out != pdf_path:
                    shutil.move(lo_out, pdf_path)
            else:
                # Fall back: drive Word COM in a dedicated STA thread
                outcome: dict = {}

                def _word_convert():
                    import pythoncom
                    import win32com.client
                    pythoncom.CoInitialize()
                    try:
                        word = win32com.client.DispatchEx("Word.Application")
                        word.Visible = False
                        opened = word.Documents.Open(docx_path)
                        opened.SaveAs(pdf_path, FileFormat=17)  # 17 = wdFormatPDF
                        opened.Close(SaveChanges=False)
                        try:
                            word.Quit()
                        except Exception:
                            pass
                    except Exception as exc:
                        outcome["error"] = str(exc)
                    finally:
                        pythoncom.CoUninitialize()

                t = threading.Thread(target=_word_convert, daemon=True)
                t.start()
                t.join(timeout=90)

                if "error" in outcome:
                    return Response(
                        {"detail": f"Conversion PDF impossible: {outcome['error']}"},
                        status=status.HTTP_500_INTERNAL_SERVER_ERROR,
                    )

            if not os.path.isfile(pdf_path):
                return Response(
                    {"detail": "Le PDF n'a pas pu être généré."},
                    status=status.HTTP_500_INTERNAL_SERVER_ERROR,
                )

            with open(pdf_path, "rb") as f:
                pdf_bytes = f.read()

        finally:
            # Clean up temp dir — ignore errors if Word still holds a lock
            shutil.rmtree(tmp_dir, ignore_errors=True)

        response = HttpResponse(pdf_bytes, content_type="application/pdf")
        response["Content-Disposition"] = (
            f'attachment; filename="decharge-{demande.id_demande}.pdf"'
        )
        return response

    @action(detail=True, methods=["post"], url_path="creer_commande_interne")
    def creer_commande_interne(self, request, pk=None):
        from apps.procurement.models import LotArticle, MarcheBC

        demande = self.get_object()

        if demande.statut not in ("partielle", "totale"):
            return Response(
                {"detail": "La demande doit etre partielle ou totale pour creer une commande interne."},
                status=status.HTTP_400_BAD_REQUEST,
            )

        if getattr(demande, "commande_interne", None):
            commande = demande.commande_interne
            return Response(
                {
                    "detail": "Une commande interne existe déjà pour cette demande.",
                    "id_marche": commande.id_marche,
                    "reference": commande.reference,
                },
                status=status.HTTP_400_BAD_REQUEST,
            )

        reference = request.data.get("reference") or f"BC-INT-{timezone.now().strftime('%Y%m%d%H%M%S%f')}"
        reference = str(reference)[:100]

        commande = MarcheBC.objects.create(
            reference=reference,
            type_acquisition="bon_commande",
            statut="en_attente_livraison",
            id_cree_par=request.user,
            id_demande_source=demande,
            beneficiaire_commande=demande.beneficiaire_nom or "",
            statut_signature_commande="non_signe",
        )

        lots = []
        for idx, ligne in enumerate(demande.lignes.select_related("id_ressource").all(), start=1):
            quantite = ligne.quantite_accordee if ligne.quantite_accordee > 0 else ligne.quantite_demandee
            lots.append(
                LotArticle(
                    numero_lot=idx,
                    designation=ligne.id_ressource.designation if ligne.id_ressource else f"Article {idx}",
                    quantite_commandee=quantite,
                    quantite_recue=0,
                    observation=ligne.observation or "",
                    id_marche=commande,
                    id_ressource=ligne.id_ressource,
                )
            )

        if lots:
            LotArticle.objects.bulk_create(lots)

        return Response(
            {
                "id_marche": commande.id_marche,
                "reference": commande.reference,
                "statut_signature_commande": commande.statut_signature_commande,
            },
            status=status.HTTP_201_CREATED,
        )

    @action(detail=True, methods=["post"], url_path="signer_commande_interne")
    def signer_commande_interne(self, request, pk=None):
        demande = self.get_object()
        commande = getattr(demande, "commande_interne", None)

        if not commande:
            return Response(
                {"detail": "Aucune commande interne liée à cette demande."},
                status=status.HTTP_404_NOT_FOUND,
            )

        if commande.statut_signature_commande == "signe":
            return Response(
                {"detail": "La commande interne est déjà signée."},
                status=status.HTTP_400_BAD_REQUEST,
            )

        commande.statut_signature_commande = "signe"
        commande.date_signature_commande = timezone.now()
        commande.save(update_fields=["statut_signature_commande", "date_signature_commande"])

        return Response(
            {
                "id_marche": commande.id_marche,
                "reference": commande.reference,
                "statut_signature_commande": commande.statut_signature_commande,
                "date_signature_commande": commande.date_signature_commande,
            },
            status=status.HTTP_200_OK,
        )

    @action(detail=True, methods=["post"], url_path="refuser")
    def refuser(self, request, pk=None):
        demande = self.get_object()

        if demande.statut not in ("en_cours",):
            return Response(
                {"detail": f"Impossible de refuser une demande au statut '{demande.statut}'."},
                status=status.HTTP_400_BAD_REQUEST,
            )

        commentaire = request.data.get("commentaire_validation", "")

        demande.statut = "refusee"
        demande.commentaire_validation = commentaire
        demande.save(update_fields=["statut", "commentaire_validation"])

        self._notify_chef(
            demande,
            titre="Demande refusée",
            message=(
                f"Votre demande #{demande.id_demande} a été refusée. "
                + (f"Motif : {commentaire}" if commentaire else "")
            ),
        )

        return Response(DemandeSerializer(demande).data, status=status.HTTP_200_OK)
