"""
Generate placeholder .docx templates for decharge PDF generation.

Run once:
    python manage.py create_decharge_templates

Then replace the generated files with your official FMPDF Word templates,
keeping the same Jinja2 variable names.

Variables used
--------------
Both templates:
    {{ date }}        – date de la demande (DD/MM/YYYY)
    {{ reference }}   – référence (e.g. DEM-0001)
    {{ service }}     – nom du service bénéficiaire

Bien inventaire rows (iterate with {%tr for ligne in lignes %} ... {%tr endfor %}):
    {{ ligne.designation }}
    {{ ligne.n_inventaire }}
    {{ ligne.qte }}
    {{ ligne.affectation }}

Consommable rows:
    {{ ligne.article }}
    {{ ligne.quantite }}
"""

from pathlib import Path

from django.core.management.base import BaseCommand
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

TEMPLATES_DIR = Path(__file__).resolve().parents[3] / "templates"


def _bold_run(paragraph, text):
    run = paragraph.add_run(text)
    run.bold = True
    return run


def _set_cell_text(cell, text, bold=False):
    cell.paragraphs[0].clear()
    run = cell.paragraphs[0].add_run(text)
    run.bold = bold


def _apply_margins(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


def _add_info_block(doc):
    tbl = doc.add_table(rows=3, cols=2)
    tbl.cell(0, 0).text = "Date :"
    tbl.cell(0, 1).text = "{{ date }}"
    tbl.cell(1, 0).text = "Référence :"
    tbl.cell(1, 1).text = "{{ reference }}"
    tbl.cell(2, 0).text = "Service bénéficiaire :"
    tbl.cell(2, 1).text = "{{ service }}"
    doc.add_paragraph()


def create_bien_inventaire(path: Path) -> None:
    doc = Document()
    _apply_margins(doc)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _bold_run(title, "FMPDF — FICHE DE DÉCHARGE\nBIENS INVENTAIRES")
    doc.add_paragraph()

    _add_info_block(doc)

    # Articles table
    # Row 0 : header
    # Row 1 : docxtpl loop row (repeated once per ligne)
    tbl = doc.add_table(rows=2, cols=4)
    tbl.style = "Table Grid"

    headers = ["DESIGNATION DU MATERIEL", "N° INVENTAIRE", "QTE", "AFFECTATION"]
    for col, h in enumerate(headers):
        _set_cell_text(tbl.cell(0, col), h, bold=True)

    # Loop row — docxtpl repeats this entire <w:tr> for every ligne
    data = tbl.rows[1].cells
    data[0].paragraphs[0].add_run("{%tr for ligne in lignes %}{{ ligne.designation }}")
    data[1].paragraphs[0].add_run("{{ ligne.n_inventaire }}")
    data[2].paragraphs[0].add_run("{{ ligne.qte }}")
    data[3].paragraphs[0].add_run("{{ ligne.affectation }}{%tr endfor %}")

    doc.add_paragraph()
    doc.add_paragraph("Signature du gestionnaire : ______________________________")
    doc.save(str(path))


def create_consommable(path: Path) -> None:
    doc = Document()
    _apply_margins(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _bold_run(title, "FMPDF — FICHE DE DÉCHARGE\nCONSOMMBLES")
    doc.add_paragraph()

    _add_info_block(doc)

    tbl = doc.add_table(rows=2, cols=2)
    tbl.style = "Table Grid"

    _set_cell_text(tbl.cell(0, 0), "ARTICLE", bold=True)
    _set_cell_text(tbl.cell(0, 1), "QUANTITE", bold=True)

    data = tbl.rows[1].cells
    data[0].paragraphs[0].add_run("{%tr for ligne in lignes %}{{ ligne.article }}")
    data[1].paragraphs[0].add_run("{{ ligne.quantite }}{%tr endfor %}")

    doc.add_paragraph()
    doc.add_paragraph("Signature du gestionnaire : ______________________________")
    doc.save(str(path))


class Command(BaseCommand):
    help = "Create placeholder .docx templates used by the /imprimer endpoint."

    def handle(self, *args, **options):
        TEMPLATES_DIR.mkdir(parents=True, exist_ok=True)

        bi_path = TEMPLATES_DIR / "decharge_bien_inventaire.docx"
        co_path = TEMPLATES_DIR / "decharge_consommable.docx"

        create_bien_inventaire(bi_path)
        self.stdout.write(self.style.SUCCESS(f"Created: {bi_path}"))

        create_consommable(co_path)
        self.stdout.write(self.style.SUCCESS(f"Created: {co_path}"))

        self.stdout.write(
            self.style.WARNING(
                "\nReplace these placeholders with your official FMPDF .docx files.\n"
                "Keep the same Jinja2 variable names ({{ date }}, {{ service }}, etc.)."
            )
        )
